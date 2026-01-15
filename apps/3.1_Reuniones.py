import streamlit as st
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
import pandas as pd
import io
import json
import re

# --- 1. CONFIGURACIÓN DE PÁGINA ---
# (Opcional, si da error puedes comentarla)
# st.set_page_config(page_title="Planificador de Reuniones", page_icon="📅")

# --- 2. CONEXIÓN CON LA IA ---
try:
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
except Exception:
    st.error("⚠️ Error: No se encontró la API KEY. Revisa el archivo .streamlit/secrets.toml")
    st.stop()

# --- 3. FUNCIONES LÓGICAS ---
def generar_planificacion(tema, objetivo, duracion):
    try:
        # CORRECCIÓN: Usamos el modelo correcto "gemini-1.5-flash"
        model = genai.GenerativeModel("models/gemini-2.5-flash")
        
        prompt = f"""
        Actúa como un Facilitador Experto. Diseña una agenda para una reunión de {duracion} minutos.
        TEMA: {tema} | OBJETIVO: {objetivo}
        
        Responde EXCLUSIVAMENTE con un JSON válido.
        Estructura:
        {{
            "agenda": [
                {{"minutos": "00-05", "actividad": "Actividad breve", "responsable": "Rol o Nombre"}},
                {{"minutos": "...", "actividad": "...", "responsable": "..."}}
            ],
            "consejos": "Consejo práctico 1... Consejo práctico 2..."
        }}
        """
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return str(e)

def procesar_respuesta(texto_completo):
    try:
        # Limpieza robusta: Busca el primer "{" y el último "}" para aislar el JSON
        texto_limpio = texto_completo.replace("```json", "").replace("```", "").strip()
        
        # Si la IA puso texto antes o después, lo recortamos
        inicio = texto_limpio.find("{")
        fin = texto_limpio.rfind("}") + 1
        if inicio != -1 and fin != 0:
            texto_limpio = texto_limpio[inicio:fin]

        datos = json.loads(texto_limpio)
        return datos.get("agenda", []), datos.get("consejos", "Sin consejos")
    except Exception as e:
        return None, f"Error interpretando la respuesta de la IA. Intenta de nuevo. ({e})"

# --- 4. FUNCIONES DE EXPORTACIÓN ---

def crear_word(tema, objetivo, agenda_lista, consejos):
    doc = Document()
    doc.add_heading(f'Plan de Reunión: {tema}', 0)
    doc.add_paragraph(f"Objetivo: {objetivo}")
    
    doc.add_heading('Agenda', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Minutos'
    hdr[1].text = 'Actividad'
    hdr[2].text = 'Responsable'
    
    for item in agenda_lista:
        row = table.add_row().cells
        row[0].text = str(item.get("minutos", ""))
        row[1].text = str(item.get("actividad", ""))
        row[2].text = str(item.get("responsable", ""))
    
    doc.add_heading('Consejos', level=1)
    doc.add_paragraph(consejos)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

def crear_pdf(tema, objetivo, agenda_lista, consejos):
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 14)
            self.cell(0, 10, 'Planificador de Reuniones', 0, 1, 'C')
            self.ln(5)

    pdf = PDF()
    pdf.add_page()
    # Función auxiliar para caracteres latinos
    def L(t): return t.encode('latin-1', 'replace').decode('latin-1') if t else ""

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, L(f"Tema: {tema}"), 0, 1)
    pdf.set_font("Arial", '', 11)
    pdf.multi_cell(0, 6, L(f"Objetivo: {objetivo}"))
    pdf.ln(5)
    
    # Tabla
    pdf.set_fill_color(240, 240, 240)
    pdf.set_font("Arial", 'B', 10)
    pdf.cell(30, 10, "Minutos", 1, 0, 'C', 1)
    pdf.cell(110, 10, "Actividad", 1, 0, 'C', 1)
    pdf.cell(50, 10, "Responsable", 1, 1, 'C', 1)
    
    pdf.set_font("Arial", '', 10)
    for item in agenda_lista:
        pdf.cell(30, 10, L(str(item.get("minutos", ""))), 1)
        pdf.cell(110, 10, L(str(item.get("actividad", ""))), 1)
        pdf.cell(50, 10, L(str(item.get("responsable", ""))), 1, 1)
        
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, "Consejos:", 0, 1)
    pdf.set_font("Arial", '', 11)
    pdf.multi_cell(0, 6, L(consejos))
    
    return pdf.output(dest='S').encode('latin-1'), "application/pdf"

def crear_excel(tema, objetivo, agenda_lista, consejos):
    # Creamos un DataFrame con los datos de la agenda
    df = pd.DataFrame(agenda_lista)
    
    output = io.BytesIO()
    # Usamos ExcelWriter para guardar
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Agenda')
        
        # Guardamos los consejos y objetivos en otra hoja
        info_extra = pd.DataFrame([
            {"Dato": "Tema", "Valor": tema},
            {"Dato": "Objetivo", "Valor": objetivo},
            {"Dato": "Consejos", "Valor": consejos}
        ])
        info_extra.to_excel(writer, index=False, sheet_name='Detalles')
        
    return output.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

# --- 5. INTERFAZ ---
st.header("Planificador de Reuniones")
st.write("Planifica tus reuniones en segundos, asignando tiempo intervención, temas a tratar y responsables.")
st.divider()

col1, col2 = st.columns([2, 1])
with col1:
    tema_input = st.text_input("Tema:", placeholder="Ej: Planificación Q1")
    obj_input = st.text_input("Objetivo:", placeholder="Ej: Aprobar presupuesto")
with col2:
    duracion_input = st.selectbox("Minutos:", [15, 30, 45, 60], index=1)

if 'resultado_agenda' not in st.session_state:
    st.session_state.resultado_agenda = None
    st.session_state.consejos_agenda = None

if st.button("⚡ Generar Planificación", type="primary", use_container_width=True):
    if not tema_input or not obj_input:
        st.warning("⚠️ Completa los campos.")
    else:
        with st.spinner("Creando estrategia..."):
            texto_raw = generar_planificacion(tema_input, obj_input, duracion_input)
            agenda_data, consejos_data = procesar_respuesta(texto_raw)
            
            if agenda_data:
                st.session_state.resultado_agenda = agenda_data
                st.session_state.consejos_agenda = consejos_data
            else:
                st.error(consejos_data)

# RESULTADOS Y DESCARGA
if st.session_state.resultado_agenda:
    st.subheader("📋 Tu Agenda")
    # Convertimos a DataFrame para mostrarlo bonito en pantalla
    st.table(pd.DataFrame(st.session_state.resultado_agenda))
    st.info(f"**💡 Tips:** {st.session_state.consejos_agenda}")
    
    st.divider()
    st.subheader("📥 Descargar Archivo")
    
    # Configuración de descarga
    c_nombre, c_tipo = st.columns([2, 1])
    
    with c_nombre:
        nombre_archivo = st.text_input("Nombre del archivo:", value="Agenda_Reunion")
    
    with c_tipo:
        tipo_archivo = st.radio("Formato:", ["Word", "PDF", "Excel"], horizontal=True)
    
    # Botón de descarga
    if st.button("💾 Descargar ahora"):
        archivo_data = None
        mime_type = ""
        ext = ""
        
        try:
            if tipo_archivo == "Word":
                archivo_data, mime_type = crear_word(tema_input, obj_input, st.session_state.resultado_agenda, st.session_state.consejos_agenda)
                ext = "docx"
            elif tipo_archivo == "PDF":
                archivo_data, mime_type = crear_pdf(tema_input, obj_input, st.session_state.resultado_agenda, st.session_state.consejos_agenda)
                ext = "pdf"
            elif tipo_archivo == "Excel":
                archivo_data, mime_type = crear_excel(tema_input, obj_input, st.session_state.resultado_agenda, st.session_state.consejos_agenda)
                ext = "xlsx"
                
            st.download_button(
                label=f"Confirmar descarga {ext}",
                data=archivo_data,
                file_name=f"{nombre_archivo}.{ext}",
                mime=mime_type,
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Error generando el archivo: {e}")
import streamlit as st
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
import io
import re

# --- 1. CONFIGURACIÓN DE PÁGINA ---
# st.set_page_config(
#    page_title="Matriz de Delegación Situacional",
#   page_icon="🎯",
#   layout="centered"
#)

# --- 2. MODELO FIJO ---
# Definimos el modelo aquí para que sea fácil de cambiar en el futuro si es necesario
MODELO_FIJO = "models/gemma-3-27b-it"

# Estilos CSS
st.markdown("""
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .block-container {padding-top: 2rem;}
    .stAlert { margin-bottom: 1rem; }
    </style>
    """, unsafe_allow_html=True)

# --- 3. CONFIGURACIÓN API ---
try:
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
except Exception as e:
    st.error("⚠️ Error: No se encontró la API KEY en secrets.toml")
    st.stop()

# --- 4. FUNCIONES DE ARCHIVOS (WORD Y PDF) ---
def crear_word(data, tarea, colaborador):
    doc = Document()
    doc.add_heading('Plan de Delegación Situacional', 0)
    doc.add_paragraph(f"Colaborador: {colaborador}")
    doc.add_paragraph(f"Tarea: {tarea}")
    
    doc.add_heading('1. Diagnóstico', level=1)
    doc.add_paragraph(data.get('diagnostico', ''))
    
    doc.add_heading('2. Pasos Clave', level=1)
    doc.add_paragraph(data.get('pasos', ''))
    
    doc.add_heading('3. Guion de Conversación', level=1)
    doc.add_paragraph(data.get('guion', ''))
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"

def crear_pdf(data, tarea, colaborador):
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 16)
            self.cell(0, 10, 'Plan de Delegación Situacional', 0, 1, 'C')
            self.ln(10)

    pdf = PDF()
    pdf.add_page()
    
    # Función para limpiar caracteres latinos y evitar errores de codificación
    def L(t): 
        if not t: return ""
        return t.encode('latin-1', 'replace').decode('latin-1')
    
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 8, L(f"Colaborador: {colaborador}"), 0, 1)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 6, L(f"Tarea: {tarea}"))
    pdf.ln(5)
    
    secciones = [
        ("1. Diagnóstico", "diagnostico"), 
        ("2. Pasos Clave", "pasos"), 
        ("3. Guion de Conversación", "guion")
    ]
    
    for titulo, key in secciones:
        pdf.set_font("Arial", 'B', 12)
        pdf.set_text_color(0, 50, 100)
        pdf.cell(0, 10, L(titulo), 0, 1)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 6, L(data.get(key, '')))
        pdf.ln(5)
        
    return pdf.output(dest='S').encode('latin-1'), "application/pdf", "pdf"

# --- 5. LÓGICA IA ---
def analizar_delegacion(tarea, exp, disp):
    try:
        # Usamos el modelo fijo definido arriba
        model = genai.GenerativeModel(MODELO_FIJO)
        
        prompt = f"""
        Actúa como Coach experto en Liderazgo Situacional.
        Tarea: {tarea} | Nivel: {exp} | Disposición: {disp}
        
        USA EXACTAMENTE ESTOS TÍTULOS PARA SEPARAR LAS SECCIONES:
        
        SECCION_DIAGNOSTICO:
        [Identifica si es E1, E2, E3 o E4 y explica por qué]
        
        SECCION_PASOS:
        [Lista 3 pasos clave para la reunión]
        
        SECCION_GUION:
        [Escribe el diálogo exacto entre comillas]
        """
        
        response = model.generate_content(prompt)
        texto = response.text
        
        # Regex (Expresiones Regulares) para cortar el texto de forma segura
        diag = re.search(r"SECCION_DIAGNOSTICO:(.*?)(?=SECCION_PASOS:|$)", texto, re.DOTALL | re.IGNORECASE)
        pasos = re.search(r"SECCION_PASOS:(.*?)(?=SECCION_GUION:|$)", texto, re.DOTALL | re.IGNORECASE)
        guion = re.search(r"SECCION_GUION:(.*?)(?=$)", texto, re.DOTALL | re.IGNORECASE)
        
        return {
            "diagnostico": diag.group(1).strip().replace("*", "") if diag else "Error generando diagnóstico.",
            "pasos": pasos.group(1).strip().replace("*", "") if pasos else "Error generando pasos.",
            "guion": guion.group(1).strip().replace("*", "").replace('"', '') if guion else "Error generando guion."
        }

    except Exception as e:
        return {"error": str(e)}

# --- 6. INTERFAZ DE USUARIO ---
st.header("Delegación Situacional")
st.write("Delega de acuerdo al grado de Competencia y Compromiso que tenga la persona a la cual delegarás la tarea")
#st.caption(f"Motor IA: {MODELO_FIJO}") # Confirmación visual del modelo usado

st.divider()

col_u1, col_u2 = st.columns(2)
with col_u1:
    nombre_colab = st.text_input("Nombre del Colaborador:", value="Juan Pérez")
with col_u2:
    st.write("") # Espaciador

col1, col2 = st.columns(2)
with col1:
    nivel_experiencia = st.selectbox("Nivel de Competencia (Hacer)", [
        "M1 - Principiante (Baja competencia)",
        "M2 - Aprendiz (Competencia moderada)",
        "M3 - Avanzado (Competencia alta)",
        "M4 - Experto (Alta competencia)"
    ])
with col2:
    disposicion = st.selectbox("Nivel de Motivación (Querer)", [
        "Bajo (Inseguro o no dispuesto)",
        "Variable (Motivado pero sin experiencia)",
        "Variable (Motivado pero cauteloso)",
        "Alto (Seguro, motivado y capaz)"
    ])

tarea = st.text_area("Tarea a delegar", height=100, placeholder="Ej: Realizar el informe mensual...")

# Inicialización de estado
if 'resultado_del' not in st.session_state:
    st.session_state.resultado_del = None

# Botón de acción
if st.button("🚀 Generar Estrategia", type="primary", use_container_width=True):
    if not tarea:
        st.warning("⚠️ Escribe una tarea primero.")
    else:
        with st.spinner("Analizando situación..."):
            st.session_state.resultado_del = analizar_delegacion(
                tarea, nivel_experiencia, disposicion
            )

# --- 7. MOSTRAR RESULTADOS ---
if st.session_state.resultado_del:
    res = st.session_state.resultado_del
    
    if "error" in res:
        st.error(f"Error técnico: {res['error']}")
    else:
        # Asignamos variables antes de imprimir para evitar errores de sintaxis en f-strings complejos
        txt_diagnostico = res.get('diagnostico', 'Sin datos')
        txt_pasos = res.get('pasos', 'Sin datos')
        txt_guion = res.get('guion', 'Sin datos')

        st.success(f"**Diagnóstico:**\n\n{txt_diagnostico}")
        st.info(f"**Pasos Clave:**\n\n{txt_pasos}")
        
        st.markdown("### 🗣️ Guion Sugerido")
        with st.container(border=True):
            st.markdown(txt_guion)
        
        st.divider()
        
        # --- ZONA DE DESCARGA ---
        st.subheader("📥 Descargar Plan")
        c_name, c_type = st.columns([2, 1])
        with c_name:
            f_name = st.text_input("Nombre del archivo:", value=f"Delegacion_{nombre_colab.split()[0]}")
        with c_type:
            f_fmt = st.radio("Formato:", ["Word (.docx)", "PDF (.pdf)"], horizontal=True)
            
        # Generación del archivo seleccionado
        if f_fmt == "Word (.docx)":
            b_data, b_mime, b_ext = crear_word(res, tarea, nombre_colab)
        else:
            b_data, b_mime, b_ext = crear_pdf(res, tarea, nombre_colab)
            
        st.download_button(
            label=f"💾 Bajar {f_fmt}",
            data=b_data,
            file_name=f"{f_name}.{b_ext}",
            mime=b_mime,
            use_container_width=True
        )
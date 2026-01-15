import streamlit as st
import google.generativeai as genai
from docx import Document
import io

# --- 1. CONFIGURACIÓN ---
# Nota: Si este archivo se ejecuta desde Inicio.py, esta línea podría ser ignorada.
# st.set_page_config(page_title="Negociador Método Harvard", page_icon="🤝", layout="centered")

# --- 2. CONEXIÓN IA ---
try:
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
except Exception:
    st.error("⚠️ Falta API Key. Asegúrate de tener el archivo .streamlit/secrets.toml configurado.")
    st.stop()

# --- 3. LÓGICA HARVARD ---
def analizar_negociacion(rol, contraparte, problema, intereses_mios, intereses_ellos, maan):
    try:
        # CORRECCIÓN: Usamos el modelo correcto
        model = genai.GenerativeModel("models/gemini-2.5-flash")
        
        prompt = f"""
        Actúa como un Experto en Negociación del 'Harvard Negotiation Project' (Fisher & Ury).
        Tu cliente es un novato que necesita una guía paso a paso.
        
        CONTEXTO:
        - Usuario: {rol}
        - Contraparte: {contraparte}
        - Conflicto: {problema}
        - Intereses del Usuario (Subyacentes): {intereses_mios}
        - Intereses de la Contraparte (Estimados): {intereses_ellos}
        - MAAN (Plan B si no hay acuerdo): {maan}
        
        TAREA: Genera una hoja de ruta estratégica.
        
        FORMATO DE RESPUESTA (Usa Títulos en Negrita):
        
        1. DIAGNÓSTICO DE PODER
        Analiza el MAAN del usuario. ¿Es fuerte o débil? ¿Debe revelarlo o mejorarlo?
        
        2. ESTRATEGIA A: CREACIÓN DE VALOR (Ideal)
        Diseña una propuesta que satisfaga los intereses de ambos (Opciones de Mutuo Beneficio).
        - Propuesta creativa: [Detalle]
        - Frase de apertura ("Speech"): [Escribe el guion exacto]
        
        3. ESTRATEGIA B: CRITERIOS OBJETIVOS (Defensiva)
        Si la contraparte se pone dura o regatea por posición.
        - Qué estándar independiente usar (precios de mercado, leyes, precedentes).
        - Frase para re-encuadrar: "No hablemos de lo que tú quieres o yo quiero, veamos qué es lo justo basado en..."
        
        4. PREGUNTAS PODEROSAS
        3 preguntas que el usuario debe hacer para descubrir más información en la mesa.
        """
        
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

# --- ESTA ES LA FUNCIÓN QUE FALTABA ---
def crear_docx(texto):
    doc = Document()
    doc.add_heading('Plan de Negociación Harvard', 0)
    
    # Limpieza básica para el Word
    lines = texto.split('\n')
    for line in lines:
        # Detectamos si es un título para darle formato
        if "DIAGNÓSTICO" in line or "ESTRATEGIA" in line or "PREGUNTAS" in line:
            doc.add_heading(line.replace('*', ''), level=1)
        else:
            doc.add_paragraph(line.replace('*', ''))
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 4. INTERFAZ ---
st.header("El Negociador Harvard")
st.write("Negocia por intereses (Personas) y no por posición (Problema).")

with st.expander("📚 ¿Qué es el Método Harvard? (Leer antes de empezar)"):
    st.info("""
    Este método no busca 'ganar' aplastando al otro, sino lograr un acuerdo sensato.
    1. **Intereses:** No te enfoques en 'posiciones' (quiero $100), sino en 'intereses' (necesito pagar la renta).
    2. **Opciones:** Busca soluciones creativas donde ambos ganen algo.
    3. **Criterios:** Usa datos objetivos (mercado, ley) para decidir, no la voluntad.
    4. **MAAN:** Tu 'As bajo la manga'. Tu plan B si no hay acuerdo.
    """)

st.divider()

col1, col2 = st.columns(2)
with col1:
    rol = st.text_input("Tu Rol:", placeholder="Ej: Proveedor de Servicios")
    intereses_mios = st.text_area("Tus Intereses (¿Para qué quieres esto?):", placeholder="Más allá del dinero/posición. Ej: Quiero estabilidad, prestigio, tiempo libre...")

with col2:
    contraparte = st.text_input("La Contraparte:", placeholder="Ej: Gerente de Compras")
    intereses_ellos = st.text_area("Intereses de Ellos (¿Qué les preocupa?):", placeholder="Ej: No pasarse del presupuesto, quedar bien con su jefe, rapidez...")

problema = st.text_area("El Conflicto/Tema a negociar:", placeholder="Ej: Renovación de contrato con aumento de tarifas del 20%.")

st.header("Tu Poder (MAAN)")
st.caption("MAAN = Mejor Alternativa al Acuerdo Negociado. Es tu Plan B real.")
maan = st.text_input("Si NO llegas a un acuerdo, ¿qué harás?", placeholder="Ej: Tengo otra oferta lista de la empresa X / Me quedo sin cliente.")

# Inicializar estado
if 'resultado_negociacion' not in st.session_state:
    st.session_state.resultado_negociacion = None

if st.button("🧠 Generar Estrategias", type="primary"):
    if not intereses_mios or not maan:
        st.warning("⚠️ Para Harvard, es crucial definir tus Intereses y tu MAAN.")
    else:
        with st.spinner("Analizando intereses y opciones de mutuo beneficio..."):
            estrategia = analizar_negociacion(rol, contraparte, problema, intereses_mios, intereses_ellos, maan)
            st.session_state.resultado_negociacion = estrategia

# --- 5. RESULTADOS ---
if st.session_state.resultado_negociacion:
    res = st.session_state.resultado_negociacion
    
    st.divider()
    st.subheader("📋 Hoja de Ruta")
    st.markdown(res)
    
    # Descarga
    docx = crear_docx(res)
    st.download_button(
        label="📥 Descargar Plan (.docx)",
        data=docx,
        file_name="Estrategia_Harvard.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
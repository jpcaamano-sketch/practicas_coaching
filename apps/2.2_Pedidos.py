import streamlit as st
import os
import io
import google.generativeai as genai
from docx import Document
from dotenv import load_dotenv

# --- 1. CONFIGURACIÓN DE PÁGINA ---
#st.set_page_config(page_title="Pedidos Impecables", page_icon="🗣️", layout="centered")

# --- 2. ESTILOS CSS ---
#st.markdown("""
#   <style>
#    .stTextArea textarea { font-size: 16px !important; }
#    .stTextInput input { font-size: 16px !important; }
#    .info-box { background-color: #f0f8ff; padding: 15px; border-radius: 10px; border-left: 5px solid #1f77b4; }
#    #MainMenu {visibility: hidden;}
#    header {visibility: hidden;}
#    footer {visibility: hidden;}
#    </style>
#    """, unsafe_allow_html=True)

# --- 3. CONEXIÓN IA (SEGURA) ---
try:
    # Intenta leer de los secretos (local o nube)
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
except Exception:
    st.error("⚠️ No se encontró la clave en .streamlit/secrets.toml")
    st.stop()

# --- 4. FUNCIONES LÓGICAS ---

def generar_pedido_ia(oyente, accion, condiciones, tiempo, contexto):
    """Genera el texto usando Google Gemini"""
    try:
        # Usamos gemini-1.5-flash para asegurar compatibilidad
        model = genai.GenerativeModel("gemini-2.5-flash")
        
        prompt = f"""
        Actúa como un Coach Ontológico experto en Fernando Flores.
        Redacta un "PEDIDO IMPECABLE" basado en:
        
        1. OYENTE: {oyente}
        2. ACCIÓN: {accion}
        3. CONDICIONES DE SATISFACCIÓN: {condiciones}
        4. TIEMPO: {tiempo}
        5. TRASFONDO: {contexto}

        Genera dos partes separadas claramente por la etiqueta "SECCION_ANALISIS":
        
        Parte 1: El GUION (listo para copiar/pegar, tono profesional y asertivo).
        Parte 2: SECCION_ANALISIS: Una explicación breve de por qué este pedido reduce incertidumbre.
        """
        
        response = model.generate_content(prompt)
        text = response.text
        
        if "SECCION_ANALISIS" in text:
            parts = text.split("SECCION_ANALISIS")
            guion = parts[0].replace("SECCION_ANALISIS", "").replace("Parte 1:", "").strip()
            analisis = parts[1].replace("Parte 2:", "").strip()
        else:
            guion = text
            analisis = "No se generó el análisis detallado."
            
        return guion, analisis
    except Exception as e:
        return f"Error al generar: {e}", ""

def crear_docx(guion, analisis):
    """Crea el archivo Word descargable"""
    doc = Document()
    doc.add_heading('PEDIDO IMPECABLE', 0)
    
    doc.add_heading('Guion Sugerido:', level=1)
    doc.add_paragraph(guion)
    
    doc.add_heading('Análisis Ontológico:', level=1)
    doc.add_paragraph(analisis)
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 5. INTERFAZ DE USUARIO ---

st.header("Pedidos Impecables")
# st.caption("Basado en la Ontología del Lenguaje")
st.write("Ahorra tiempo, recursos, reprocesos, etc., haciendo Pedidos Efectivos a la primera")

with st.container(border=True):
    col1, col2 = st.columns(2)
    with col1:
        oyente = st.text_input("1. ¿A quién le pides?", placeholder="Ej: Juan, Jefe de Marketing")
    with col2:
        tiempo = st.text_input("2. ¿Para cuándo?", placeholder="Ej: Viernes 20 antes de las 14:00")

    accion = st.text_area("3. ¿Qué necesitas que se realice en el futuro?", placeholder="Ej: Que envíes el reporte de ventas", height=100)
    condiciones = st.text_area("4. ¿Qué condiciones se deben cumplir para que quedes satisfecho con el resultado?", placeholder="Ej: Formato PDF, incluyendo gráficos de Q1", height=100)
    contexto = st.text_area("5. ¿Para qué necesitas lo que pides? (Le indica al otro la importancia para qué lo necesitas)", placeholder="Ej: Para la reunión de directorio del lunes", height=100)

    # Botón de acción
    if st.button("🚀 GENERAR PEDIDO", type="primary", use_container_width=True):
        if not oyente or not accion or not tiempo:
            st.warning("⚠️ Faltan datos clave: Oyente, Acción y Tiempo son obligatorios.")
        else:
            with st.spinner("Construyendo acto del habla con IA..."):
                guion_gen, analisis_gen = generar_pedido_ia(oyente, accion, condiciones, tiempo, contexto)
                
                # Verificación de error en la respuesta
                if "Error al generar" in guion_gen:
                    st.error(guion_gen)
                else:
                    st.session_state.pedido = {"guion": guion_gen, "analisis": analisis_gen, "oyente": oyente}
                    st.rerun()

# --- 6. RESULTADOS ---
if 'pedido' in st.session_state:
    res = st.session_state.pedido
    
    st.divider()
    st.subheader("📄 TU PEDIDO LISTO:")
    
    # Visualización
    st.markdown(f"""
    <div class="info-box">
        <b>Guion Sugerido:</b><br><br>
        {res['guion'].replace(chr(10), '<br>')}
    </div>
    """, unsafe_allow_html=True)
    
    with st.expander("🧠 Ver Análisis (Por qué funciona)"):
        st.write(res['analisis'])
    
    col_d1, col_d2 = st.columns(2)
    
    # Descargar TXT (Simple)
    with col_d1:
        st.download_button(
            label="💾 Descargar Texto (.txt)",
            data=res['guion'],
            file_name=f"pedido_{res['oyente']}.txt",
            mime="text/plain",
            use_container_width=True
        )
    
    # Descargar DOCX (Completo)
    with col_d2:
        docx_file = crear_docx(res['guion'], res['analisis'])
        st.download_button(
            label="📝 Descargar Word (.docx)",
            data=docx_file,
            file_name=f"pedido_{res['oyente']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    # Botón reiniciar
    if st.button("🔄 Hacer Nuevo Pedido", use_container_width=True):
        del st.session_state.pedido
        st.rerun()